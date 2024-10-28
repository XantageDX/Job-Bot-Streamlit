import streamlit as st
import openai
import bcrypt
from io import BytesIO
from docx import Document

# Load API key and user credentials from secrets
openai_api_key = st.secrets["OPENAI_API_KEY"]
user_credentials = st.secrets["users"]

# Ensure the OpenAI API key is present
if not openai_api_key:
    st.error("OPENAI_API_KEY is missing. Please ensure it is set in the secrets.toml file.")
else:
    client = openai.OpenAI(api_key=openai_api_key)

# Password verification function
def check_password(plain_password, hashed_password):
    """Verifies if the entered password matches the stored hashed password."""
    return bcrypt.checkpw(plain_password.encode(), hashed_password.encode())

# Sidebar login function
def login():
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False  # Default to not authenticated

    if not st.session_state["authenticated"]:
        # Sidebar content
        st.sidebar.image("xantage.jpg", caption="Powered by Xantage", use_column_width=True)
        
        with st.sidebar.form("login_form"):
            st.write("### Log In")
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submit = st.form_submit_button("Log In")

            if submit:
                if username in user_credentials:
                    hashed_password = user_credentials[username]["password"]  # Retrieve hashed password

                    if check_password(password, hashed_password):
                        st.session_state["authenticated"] = True
                        st.session_state["username"] = username  # Store username in session state
                        st.sidebar.success(f"Welcome, {username}!")
                    else:
                        st.sidebar.error("Incorrect password.")
                else:
                    st.sidebar.error("User not found.")
    else:
        st.sidebar.image("xantage.jpg", caption="Powered by Xantage", use_column_width=True)
        st.sidebar.success(f"Welcome, {st.session_state['username']}!")
        st.sidebar.button("Log Out", on_click=lambda: st.session_state.update(authenticated=False, username=""))

# Display login form in the sidebar and check authentication
login()
if not st.session_state.get("authenticated"):
    st.stop()  # Stop the app here if the user is not authenticated

# Initialize session state variables for app
if 'job_ad' not in st.session_state:
    st.session_state['job_ad'] = ""

if 'file_name' not in st.session_state:
    st.session_state['file_name'] = "job_ad"

if 'download_clicked' not in st.session_state:
    st.session_state['download_clicked'] = False

# Main page content for generating a job ad
st.title("Job Ad Generator")

# Form for job details
with st.form(key='job_ad_form'):
    job_title = st.text_input("Job Title")
    about_company = st.text_area("About the Company")
    qualifications = st.text_area("Qualifications")
    years_of_experience = st.text_input("Years of Experience")
    benefits = st.text_area("Benefits")
    job_type = st.selectbox("Job Type", ["On-site", "Hybrid", "Remote"])
    contract_type = st.selectbox("Contract Type", ["Full-time", "Part-time", "Contract", "Internship", "Freelance"])
    location = st.text_input("Location")
    
    submit_button = st.form_submit_button(label='Send to AI')

if submit_button:
    # Show a spinner while waiting for the model response
    with st.spinner("Waiting for model response..."):
        prompt = f"""
        I need you to create a job ad based on the following structure:
        Title: {job_title}\n\n
        About the Company: {about_company}\n\n
        Role Overview: Create a brief overview of the role in {job_title} based on the required qualifications {qualifications}\n\n
        Key Responsibilities: Generate a bullet point list of key responsibilities for the role {job_title}\n\n
        Qualifications: Create a bullet list copying the data from {qualifications}\n\n
        Years of experience: {years_of_experience}\n\n
        Create a compelling conclusion to encourage candidates to apply\n\n
        Contract Type: {contract_type}\n\n
        Benefits: list from {benefits}\n\n
        Job Type: {job_type}\n\n
        Location: {location}
        """

        # Make a request to OpenAI's API
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000
        )

        # Store the generated job ad in session state
        st.session_state['job_ad'] = response.choices[0].message.content.strip()
        st.session_state['download_clicked'] = False

# Display the generated job ad
if st.session_state['job_ad']:
    st.write("### Generated Job Ad")
    st.write(st.session_state['job_ad'])

    # File name input for download
    st.session_state['file_name'] = st.text_input("Enter the file name (without extension)", st.session_state['file_name'])

    if st.session_state['file_name']:
        # Create a Word document
        doc = Document()
        doc.add_heading('Job Ad', 0)
        doc.add_paragraph(st.session_state['job_ad'])

        # Save the doc to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)  # Move to the beginning of the BytesIO buffer

        # Download button
        download_button = st.download_button(
            label="Download as .docx",
            data=doc_io.getvalue(),
            file_name=f"{st.session_state['file_name']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Confirm download
        if download_button:
            st.session_state['download_clicked'] = True

# Show confirmation message after download
if st.session_state['download_clicked']:
    st.success("File downloaded successfully!")
