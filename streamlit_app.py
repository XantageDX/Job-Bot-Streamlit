import streamlit as st
import openai
import streamlit_authenticator as stauth
import yaml
from yaml.loader import SafeLoader
from dotenv import load_dotenv
import os
from io import BytesIO
from docx import Document

# Load environment variables from .env
#load_dotenv()

# Authentication setup
with open('config.yaml') as file:
    config = yaml.load(file, Loader=SafeLoader)

authenticator = stauth.Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days']
)

# Initialize session state variables
if 'job_ad' not in st.session_state:
    st.session_state['job_ad'] = ""

if 'file_name' not in st.session_state:
    st.session_state['file_name'] = "job_ad"

if 'download_clicked' not in st.session_state:
    st.session_state['download_clicked'] = False

name, authentication_status, username = authenticator.login('Login form', location='sidebar')

if authentication_status:
    # Sidebar content
    st.sidebar.image("xantage.jpg", caption="Powered by Xantage", use_column_width=True)
    st.sidebar.write(f'Welcome *{name}*')
    authenticator.logout('Logout', 'sidebar')

    # Main page content
    st.title("Job Ad Generator")

    # Form for job details
    with st.form(key='job_ad_form'):
        job_title = st.text_input("Job Title")
        about_company = st.text_area("About the Company")
        qualifications = st.text_area("Qualifications")
        years_of_experience = st.text_input("Years of Experience")
        benefits = st.text_area("Benefits")
        job_type = st.selectbox("Job Type", ["On-site", "Hybrid", "Remote"])
        contract_type = st.selectbox("Contract Type", ["Full-time", "Part-time", "Contract", "Internship", "Freelance"])
        location = st.text_input("Location")
        
        submit_button = st.form_submit_button(label='Send to AI')

    if submit_button:
        # Show a spinner while waiting for the model response
        with st.spinner("Waiting for model response..."):
            # Prompt generation
            prompt = f"""
            I need you to create a job ad based on the following structure:
            Title: {job_title}\n\n
            About the Company: {about_company}\n\n
            Role Overview: Create a brief overview of the role in {job_title} based on the required qualifications {qualifications}\n\n
            Key Responsibilities: Generate a bullet point list of key responsibilities for the role {job_title}\n\n
            Qualifications: Create a bullet list copying the data from {qualifications}\n\n
            Years of experience: {years_of_experience}\n\n
            Create a compelling conclusion to encourage candidates to apply\n\n
            Contract Type: {contract_type}\n\n
            Benefits: list from {benefits}\n\n
            Job Type: {job_type}\n\n
            Location: {location}
            """


            # Accedi al segreto direttamente dal secrets manager di Streamlit Cloud
            openai_api_key = st.secrets["OPENAI_API_KEY"]
            client = openai.OpenAI(api_key=openai_api_key)
            # OpenAI API call
            #client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

            # Request to OpenAI's API
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000
            )

            # Store the generated job ad in session state
            st.session_state['job_ad'] = response.choices[0].message.content.strip()
            st.session_state['download_clicked'] = False  # Reset the download state

    # Check if there is a job ad to display
    if st.session_state['job_ad']:
        st.write("### Generated Job Ad")
        st.write(st.session_state['job_ad'])

        # File name input (stays even when the page reruns)
        st.session_state['file_name'] = st.text_input("Enter the file name (without extension)", st.session_state['file_name'])

        if st.session_state['file_name']:
            # Create a Word document
            doc = Document()
            doc.add_heading('Job Ad', 0)
            doc.add_paragraph(st.session_state['job_ad'])

            # Save the doc to a BytesIO object
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)  # Move to the beginning of the BytesIO buffer

            # Download button with dynamic file name
            download_button = st.download_button(
                label="Download as .docx",
                data=doc_io.getvalue(),
                file_name=f"{st.session_state['file_name']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Set download flag after successful download
            if download_button:
                st.session_state['download_clicked'] = True

    # Show confirmation message after download
    if st.session_state['download_clicked']:
        st.success("File downloaded successfully!")

elif authentication_status == False:
    st.error('Username/password is incorrect')
elif authentication_status == None:
    st.warning('Please enter your username and password')
